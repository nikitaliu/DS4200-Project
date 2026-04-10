"""
Generate the DS4200 design document as a Word file.
"""

from pathlib import Path

from docx import Document

from data_cleaning import ROOT


OUTPUT_PATH = ROOT / "design_document.docx"
PUBLISHED_URL = "https://nikitaliu.github.io/DS4200-Project/"


def add_viz_section(document: Document, title: str, body: str) -> None:
    document.add_heading(title, level=2)
    document.add_paragraph(body)


def main() -> None:
    document = Document()
    document.add_heading("Massachusetts Housing Financial Analysis Dashboard", level=0)
    document.add_paragraph(
        "Authors: Xinyue Du, Xiqiao Liu, Lin Pan\n"
        "Course: DS4200 — Information Presentation and Data Visualization\n"
        "Date: April 2026"
    )

    document.add_heading("Project Summary", level=1)
    document.add_paragraph(
        "This design document explains the rationale for each visualization used in the "
        "Massachusetts Housing Financial Analysis Dashboard. The site combines D3 and Altair "
        "to connect housing prices with property characteristics, livability, affordability, "
        "and environmental risk in a format designed for a general public audience. The published "
        "website target is: "
        f"{PUBLISHED_URL}"
    )

    add_viz_section(
        document,
        "Visualization 1 — D3 Choropleth Map of Town-Level Market Metrics",
        "The choropleth was chosen because the first analytical task is geographic screening: users "
        "need to see how valuation, affordability, cap-rate proxy, and environmental risk cluster across "
        "Massachusetts. Color is the primary encoding because it is the clearest way to compare one metric "
        "across many municipalities on a map, while tooltips carry the supporting values for median income, "
        "listing count, and price-to-income ratio. A dropdown was preferred over small multiples because it "
        "keeps the map readable on a single page and lets the user reuse the same spatial memory while changing "
        "the analytical lens. A bubble map and county-level aggregation were rejected because they either distort "
        "municipal boundaries or hide town-level variation that matters in housing analysis. Clicking a town opens "
        "a profile panel so users can move from a statewide view into local housing, demographic, and employment "
        "context without leaving the page. From a financial analysis perspective, the map acts like a market-screening "
        "dashboard that helps identify premium towns, relative bargains, and places where risk and price move together "
        "instead of offsetting each other."
    )

    add_viz_section(
        document,
        "Visualization 2 — Altair Scatter Plot of Price Drivers",
        "The price-driver scatter plot was selected because comparable sales analysis is fundamentally about "
        "seeing how listings cluster when a property attribute changes. Square footage, bedrooms, year built, "
        "and livability composite appear as interchangeable x-axis choices so users can test different valuation "
        "stories without leaving the chart. Price is placed on the y-axis because it is the dependent market "
        "outcome, property type is encoded with color to separate market segments, and environmental risk is "
        "encoded with point size so the chart can carry one more context variable without sacrificing "
        "readability. Brush selection was preferred over a static summary because it lets users create their own "
        "comparison set and immediately inspect the implied mean and median price. A regression line alone was "
        "rejected because it hides distributional spread and segment differences that matter to home shoppers."
    )

    add_viz_section(
        document,
        "Visualization 3 — Altair Grouped Bar Chart of Livability by Price Quartile",
        "The grouped bar chart was chosen to compare three related livability metrics across four price quartiles "
        "using a simple categorical layout. Quartile appears on the x-axis because the goal is to compare market "
        "tiers, average score is encoded by bar height because magnitude comparison is straightforward for grouped "
        "bars, and color differentiates walk, bike, and transit scores. A stacked bar was rejected because it would "
        "obscure direct comparison between the individual livability measures, while a line chart would suggest a "
        "continuous scale that does not exist here. The financial interpretation is a livability-premium analysis: "
        "the chart shows whether higher-priced listings consistently bundle accessibility advantages and whether that "
        "premium appears strongest in walkability, transit access, or both."
    )

    add_viz_section(
        document,
        "Visualization 4 — D3 Box Plot of Price Distribution by Property Type",
        "A box-and-whisker plot is the most appropriate design for showing how each property segment differs not only "
        "in median price but also in spread, skew, and outlier behavior. Property type is placed on the x-axis as the "
        "main segmentation variable, and listing price is placed on the y-axis so the user can compare medians and IQRs "
        "at a glance. Hover tooltips expose the exact quartiles and counts, while clickable outliers reveal individual "
        "listing details. Histograms and violin plots were considered, but the box plot was preferred because it more "
        "directly communicates risk dispersion and is easier for a mixed audience to interpret quickly. This chart "
        "supports segment comparison by showing which home types have tight distributions and which contain the widest "
        "pricing uncertainty."
    )

    add_viz_section(
        document,
        "Visualization 5 — Altair Correlation Heatmap of Financial and Risk Factors",
        "The correlation heatmap was selected because the project needs a compact factor-analysis view that makes it "
        "easy to compare many variable pairs at once. A diverging red-blue palette was used so positive and negative "
        "relationships are visually separable around zero, and tooltip interaction exposes the exact coefficient for "
        "any pair. A table of coefficients was rejected because it is slower to scan and does not highlight patterns "
        "or multicollinearity as effectively. The financial value of this chart is that it surfaces which features move "
        "most strongly with price and whether environmental risks appear priced in, ignored, or intertwined with other "
        "location advantages."
    )

    add_viz_section(
        document,
        "Visualization 6 — Altair Scatter Plot of Income Versus Price by Town",
        "The income-versus-price scatter plot was chosen to compare local purchasing power against local housing prices "
        "in a way that makes overvaluation and relative affordability visible at the town level. Median household income "
        "is on the x-axis, median listing price is on the y-axis, and price-to-income ratio is encoded with color to keep "
        "the affordability story central. The dashed reference line at four times income provides an interpretable lending-style "
        "benchmark. A choropleth alone would not show the cross-town affordability relationship as clearly, and a ranked bar chart "
        "would hide how far towns sit above or below a common threshold. This chart is the clearest way to show which towns appear "
        "expensive relative to their own economic base and which look more balanced for first-time buyers or renters comparing where "
        "they may have a better chance of staying within reach."
    )

    document.add_heading("Town Profile Interaction", level=1)
    document.add_paragraph(
        "The choropleth includes a click-through town profile panel inspired by MAPC DataCommon. We chose a slide-in "
        "panel instead of generating separate HTML pages because it preserves the single-page storytelling flow required "
        "by the course while still giving users a place-specific drill-down. The Housing tab summarizes listing counts, "
        "prices, and property-type mix from the cleaned Zillow data. The Demographics tab shows ACS income, tenure, age, "
        "and race or ethnicity context. The Employment tab uses ACS industry counts to show what kinds of work are most "
        "common in each town. This interaction matters because it helps a general audience move from statewide patterns to "
        "town-level context without losing orientation."
    )

    document.add_heading("Assumptions and Tradeoffs", level=1)
    document.add_paragraph(
        "Estimated cap rate is based on a rental proxy derived from the listing’s estimated monthly payment, multiplied "
        "by 12 and then by a 1.1 rent multiplier. This assumption was documented because it is a rough heuristic rather than "
        "a direct rent observation. Town-name normalization was also necessary to align neighborhood-style listing labels "
        "with municipal census and map boundaries."
    )

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_PATH)
    print(f"Saved design document to {OUTPUT_PATH.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
